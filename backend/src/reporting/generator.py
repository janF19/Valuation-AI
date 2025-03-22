from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
from typing import Dict, Any, Optional
from openai import OpenAI
from backend.config.settings import settings
from datetime import datetime

class ReportGenerator:
    @staticmethod
    def generate(data: Dict[str, Any]) -> Document:
        doc = Document()
        
        # Add creation timestamp to the report
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Report generated: {timestamp}")
        
        # Title
        title = doc.add_heading('Company Valuation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company Overview Section
        doc.add_heading('Company Overview', level=1)
        
        # Safely access company information with fallbacks
        company_info = data.get('financial_data', {}).get('information', {})
        if not company_info:
            company_info = {
                'company_name': 'N/A',
                'industry': 'N/A',
                'headquarters': 'N/A',
                'established': 'N/A',
                'employees': 'N/A',
                'main_activities': ['N/A']
            }
            
        overview_para = doc.add_paragraph()
        overview_para.add_run(f"Company: {company_info.get('company_name', 'N/A')}\n").bold = True
        overview_para.add_run(f"Industry: {company_info.get('industry', 'N/A')}\n")
        overview_para.add_run(f"Location: {company_info.get('headquarters', 'N/A')}\n")
        overview_para.add_run(f"Established: {company_info.get('established', 'N/A')}\n")
        overview_para.add_run(f"Employees: {company_info.get('employees', 'N/A')}\n")
        
        # Get AI-generated company description
        description = ReportGenerator._get_company_description(company_info)
        doc.add_paragraph(description)
        
        # Financial Analysis Section
        doc.add_heading('Financial Analysis', level=1)
        
        # Add key metrics with safe access
        metrics = data.get('financial_data', {}).get('income_statement', {})
        if not metrics:
            metrics = {}
            
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Table Grid'
        metrics_table.rows[0].cells[0].text = 'Key Metrics (thousands Kč)'
        metrics_table.rows[0].cells[1].text = '2023'
        
        # Define metrics with fallback values
        key_metrics = [
            ('Revenue', metrics.get('revenue_from_products_and_services_current', 'N/A')),
            ('EBIT', metrics.get('ebit_current', 'N/A')),
            ('Operating Profit', metrics.get('operating_profit_current', 'N/A'))
        ]
        
        for metric, value in key_metrics:
            row = metrics_table.add_row()
            row.cells[0].text = metric
            # Format only if the value is numeric
            row.cells[1].text = f"{value:,}" if isinstance(value, (int, float)) else str(value)
            
        # Generate Financial Health Analysis
        doc.add_heading('Financial Health Assessment', level=2)
        health_analysis = ReportGenerator._analyze_financial_health(data.get('financial_data', {}))
        doc.add_paragraph(health_analysis)
        
        # Valuation Results with safer handling
        doc.add_heading('Valuation Analysis', level=1)
        
        # Handle missing valuation data
        valuation = data.get('result_valuation')
        if not valuation:
            valuation = {
                'EV/EBIT Multiple': 'N/A',
                'EV/EBITDA Multiple': 'N/A',
                'Enterprise Value based on EBIT (Kč thousands)': 'N/A',
                'Enterprise Value based on EBITDA (Kč thousands)': 'N/A',
                'EBIT': 0,
                'EBITDA': 0
            }
        
        val_para = doc.add_paragraph()
        val_para.add_run('Enterprise Value Based on Multiples:\n').bold = True
        
        # Safe access with type checking
        ev_ebit_multiple = valuation.get('EV/EBIT Multiple', 'N/A')
        ev_ebitda_multiple = valuation.get('EV/EBITDA Multiple', 'N/A')
        ev_ebit_value = valuation.get('Enterprise Value based on EBIT (Kč thousands)', 'N/A')
        ev_ebitda_value = valuation.get('Enterprise Value based on EBITDA (Kč thousands)', 'N/A')
        
        # Format values for display
        val_para.add_run(f"• EV/EBIT Multiple ({ev_ebit_multiple}x): ")
        if isinstance(ev_ebit_value, (int, float)):
            val_para.add_run(f"{ev_ebit_value:,.0f} thousands Kč\n")
        else:
            val_para.add_run(f"{ev_ebit_value}\n")
            
        val_para.add_run(f"• EV/EBITDA Multiple ({ev_ebitda_multiple}x): ")
        if isinstance(ev_ebitda_value, (int, float)):
            val_para.add_run(f"{ev_ebitda_value:,.0f} thousands Kč\n")
        else:
            val_para.add_run(f"{ev_ebitda_value}\n")
        
        # Add EBIT/EBITDA comparison graph if data is available
        if isinstance(valuation.get('EBIT'), (int, float)) and isinstance(valuation.get('EBITDA'), (int, float)):
            ReportGenerator._add_valuation_graph(doc, valuation)
        else:
            doc.add_paragraph("Insufficient data to generate EBIT/EBITDA comparison graph.")
        
        # Generate and add conclusion
        conclusion = ReportGenerator._generate_conclusion(valuation, health_analysis)
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(conclusion)
        
        return doc

    @staticmethod
    def _analyze_financial_health(financial_data: Dict) -> str:
        """Generate financial health analysis based on the financial data"""
        # Safe access to metrics
        metrics = financial_data.get('income_statement', {})
        if not metrics:
            return "Insufficient financial data available to perform a comprehensive health analysis."
        
        # Safely get metrics with default values
        revenue = metrics.get('revenue_from_products_and_services_current', 'N/A')
        ebit = metrics.get('ebit_current', 'N/A')
        operating_profit = metrics.get('operating_profit_current', 'N/A')
        
        # Format values only if they are numeric
        revenue_str = f"{revenue:,.0f}" if isinstance(revenue, (int, float)) else str(revenue)
        ebit_str = f"{ebit:,.0f}" if isinstance(ebit, (int, float)) else str(ebit)
        operating_profit_str = f"{operating_profit:,.0f}" if isinstance(operating_profit, (int, float)) else str(operating_profit)
        
        # Check if we have enough data to perform analysis
        available_metrics = [m for m in [revenue, ebit, operating_profit] if m != 'N/A']
        if not available_metrics:
            return "No financial metrics available to perform analysis."
        
        # Create analysis using OpenAI
        try:
            client = OpenAI(api_key=settings.OPENAI_API_KEY)
            prompt = f"""
            Analyze the financial health of a company based on these metrics (all values in thousands Kč):
            - Revenue: {revenue_str}
            - EBIT: {ebit_str}
            - Operating Profit: {operating_profit_str}

            Provide a brief (2-3 sentences) analysis of the company's financial health based on these metrics.
            Focus on profitability and operational efficiency. If any metrics are unavailable (shown as N/A),
            focus on the available metrics only.
            """
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350
            )
            
            return response.choices[0].message.content
        except Exception as e:
            return f"Unable to generate financial health analysis due to an error: {str(e)}"

    @staticmethod
    def _get_company_description(company_info: Dict) -> str:
        # Check if we have the minimum required information
        if not company_info.get('company_name') or company_info.get('company_name') == 'N/A':
            return "Insufficient company information available to generate a detailed description."
        
        # Safe access to main activities
        main_activities = company_info.get('main_activities', ['N/A'])
        if not isinstance(main_activities, list):
            main_activities = ['N/A']
            
        try:
            client = OpenAI(api_key=settings.OPENAI_API_KEY)
            prompt = f"""
            Write a brief (2-3 sentences) professional description for this company:
            Company: {company_info.get('company_name', 'N/A')}
            Industry: {company_info.get('industry', 'N/A')}
            Main activities: {', '.join(main_activities)}
            Location: {company_info.get('headquarters', 'N/A')}
            Established: {company_info.get('established', 'N/A')}
            """
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350
            )
            
            return response.choices[0].message.content
        except Exception as e:
            return f"Unable to generate company description due to an error: {str(e)}"

    @staticmethod
    def _add_valuation_graph(doc: Document, valuation: Dict):
        try:
            plt.figure(figsize=(8, 4))
            metrics = ['EBIT', 'EBITDA']
            
            # Safely access EBIT and EBITDA values with defaults
            ebit_value = valuation.get('EBIT', 0)
            ebitda_value = valuation.get('EBITDA', 0)
            
            # Only proceed if we have numeric values
            if not isinstance(ebit_value, (int, float)) or not isinstance(ebitda_value, (int, float)):
                doc.add_paragraph("Insufficient data to generate EBIT/EBITDA comparison graph.")
                return
                
            values = [ebit_value, ebitda_value]
            
            plt.bar(metrics, values)
            plt.title('EBIT vs EBITDA Comparison')
            plt.ylabel('Value (thousands Kč)')
            
            # Save the plot to memory
            img_stream = io.BytesIO()
            plt.savefig(img_stream, format='png', bbox_inches='tight')
            img_stream.seek(0)
            
            # Add the image to the document
            doc.add_picture(img_stream, width=Inches(6))
            plt.close()
        except Exception as e:
            doc.add_paragraph(f"Error generating valuation graph: {str(e)}")

    @staticmethod
    def _generate_conclusion(valuation: Dict, health_analysis: str) -> str:
        # Check if we have the minimum required information
        if not valuation or not isinstance(valuation, dict):
            return "Insufficient valuation data available to generate a conclusion."
            
        # Safely access valuation metrics
        ev_ebit = valuation.get('Enterprise Value based on EBIT (Kč thousands)', 'N/A')
        ev_ebitda = valuation.get('Enterprise Value based on EBITDA (Kč thousands)', 'N/A')
        
        # Format values for display in the prompt
        ev_ebit_str = f"{ev_ebit:,.0f}" if isinstance(ev_ebit, (int, float)) else str(ev_ebit)
        ev_ebitda_str = f"{ev_ebitda:,.0f}" if isinstance(ev_ebitda, (int, float)) else str(ev_ebitda)
        
        try:
            client = OpenAI(api_key=settings.OPENAI_API_KEY)
            prompt = f"""
            Generate a brief conclusion (2-3 sentences) for a company valuation report based on:
            
            EV/EBIT Valuation: {ev_ebit_str} thousands Kč
            EV/EBITDA Valuation: {ev_ebitda_str} thousands Kč
            Financial Health Analysis: {health_analysis}
            """
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350
            )
            
            return response.choices[0].message.content
        except Exception as e:
            return f"Unable to generate conclusion due to an error: {str(e)}"